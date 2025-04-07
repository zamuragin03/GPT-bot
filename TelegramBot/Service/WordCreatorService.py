from docx import Document
from docx.shared import Pt, RGBColor, Cm
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class GOSTWordBaseDocument:
    def __init__(self, data):
        self.data = data
        self.doc = Document()
        self.set_document_style()
        self.add_page_numbering()

    def set_document_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Установка полей документа
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)  # 3 см
            section.right_margin = Cm(1.5)  # 1 см

    def add_page_numbering(self):
        # Добавление нумерации страниц
        sections = self.doc.sections
        for section in sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph(
            )
            paragraph.alignment = 1  # Центрирование текста
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run.element.append(fldChar1)
            run.element.append(instrText)
            run.element.append(fldChar2)

            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(12)

    def extract_text(self, input_string):
        allowed_tags = ['h1', 'h2', 'p']
        pattern = f"<({'|'.join(allowed_tags)})>(.*?)</\\1>"
        match = re.search(pattern, input_string, re.DOTALL)
        # Если есть совпадение, возвращаем текст внутри тега; иначе - пустая строка
        return match.group(2).strip() if match else ""

    def add_paragraph(self, text):
        p = self.doc.add_paragraph()
        p.alignment = 3  # Выравнивание по ширине
        p_format = p.paragraph_format
        p_format.first_line_indent = Cm(1.25)  # Красная строка 1.25 см
        p_format.line_spacing = 1.5
        p_format.space_after = 0
        p_format.space_before = 0
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def add_heading(self, text, level):
        self.doc.add_paragraph()
        h = self.doc.add_heading(level=level)
        h.alignment = 1  # Выравнивание по ширине
        run = h.add_run(text.upper() if level == 1 else text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = False
        self.doc.add_paragraph()

    def add_watermark(self, text):
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = 1
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(45)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def create_document(self):
        for item in self.data:
            for content in item['content']:
                if content['type'] == 'text':
                    parts = self.parse_html_content(content['text'])
                    for part in parts:
                        if '<h1>' in part:
                            self.add_heading(part.replace(
                                '<h1>', '').replace('</h1>', ''), level=1)
                        elif '<h2>' in part:
                            self.add_heading(part.replace(
                                '<h2>', '').replace('</h2>', ''), level=2)
                        elif '<p>' in part:
                            clean_text = part.replace(
                                '<p>', '').replace('</p>', '')
                            self.add_paragraph(clean_text)
        return self.doc

    def save_document(self, filename):
        self.add_watermark('Сделано с помощью @student_helpergpt_bot')
        self.doc.save(filename)


class GOSTWordDocument(GOSTWordBaseDocument):
    def create_document(self):
        for index, item in enumerate(self.data):
            if type(item) == str:
                tag_text = self.extract_text(item)
                if '<h1>' in item:
                    self.add_heading(tag_text, level=1)
                elif '<h2>' in item:
                    self.add_heading(tag_text, level=2)
                continue
            if type(item) == list:
                for p in item:
                    tag_text = self.extract_text(p)
                    self.add_paragraph(tag_text)
                continue
            if index == len(self.data) - 1 and type(item) == list:
                references = [self.extract_text(p) for p in item]
                self.add_references(references)


        return self.doc

    def add_references(self, references):
        for ref in references:
            # Удаление существующей нумерации в начале строки
            clean_ref = re.sub(r'^\d+\.\s+', '', ref)
            p = self.doc.add_paragraph(clean_ref, style='List Number')
            p.alignment = 3
            p_format = p.paragraph_format
            p_format.first_line_indent = Pt(1.25 * 28.35)

            p_format.line_spacing = 1.5
            p_format.space_after = 0
            p_format.left_indent = 0
            p_format.right_indent = 0
            p_format.space_before = 0
            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


class GOSTWordEssayDocument(GOSTWordBaseDocument):
    def create_document(self):
        for index, item in enumerate(self.data):
            print(item)
            if type(item) == str:
                tag_text = self.extract_text(item)
                if '<h1>' in item:
                    self.add_heading(tag_text, level=1)
                elif '<h2>' in item:
                    self.add_heading(tag_text, level=2)
                continue
            if type(item) == list:
                for p in item:
                    tag_text = self.extract_text(p)
                    self.add_paragraph(tag_text)
                continue
        return self.doc