from io import BytesIO
from PyPDF2 import PdfReader
import time, json, itertools, xlrd, re
from typing import Union
from openpyxl import load_workbook
from PIL import Image, ImageEnhance, ImageOps
from aiogram.types import Message, CallbackQuery
from aiogram.fsm.state import State
from DataModels.AbstractDataModel import PlanResponse
from aiogram.fsm.context import FSMContext
from datetime import datetime, timedelta
import random
import aiofiles
import docx
from Service import LocalizationService
from Service.TelegramUserSubscriptionService import TelegramUserSubscriptionService
import base64
# ИМПОРТ ДЛЯ eval
import numpy as np
import matplotlib.pyplot as plt
from aiogram.types import FSInputFile
from Config import PATH_TO_TEMP_FILES, SUBSCRIPTION_LIMITATIONS, GROUP_LINK_URL, DAILY_LIMITATIONS, PATH_TO_DOWNLOADED_FILES, PATH_TO_TEMP_WATERMARK, client
from DataModels.ChartCreatorDataModel import ChartResponse
import matplotlib
from aiogram import types, Bot
import asyncio
from aiogram.enums.parse_mode import ParseMode

matplotlib.use('Agg')

matplotlib.rc('text.latex', preamble=r'\usepackage{amsmath}')


class BotService:
    progress_indicators = ["◒", "◐", "◓", "◑"]  # Разные символы вращения
    progress_cycle = itertools.cycle(progress_indicators)

    @staticmethod
    async def encode_image(message: types.Message):
        photo = message.photo[-1]
        photo_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        photo_folder.mkdir(parents=True, exist_ok=True)
        photo_path = photo_folder.joinpath(f'{photo.file_unique_id}.jpg')
        await message.bot.download(file=message.photo[-1].file_id, destination=photo_path)
        async with aiofiles.open(photo_path, "rb") as image_file:
            return base64.b64encode(await image_file.read()).decode('utf-8')

    @staticmethod
    def invert_image(base_image_path, output_path):
        """
        Инвертирует изображение (фон -> черный, текст -> белый).
        """
        # Открываем изображение
        base_image = Image.open(base_image_path).convert("RGB")

        # Инвертируем цвета изображения
        inverted_image = ImageOps.invert(base_image)

        # Сохраняем инвертированное изображение
        inverted_image.save(output_path, "JPEG")

    @staticmethod
    def latex_to_image(latex_code, external_id, dpi=300):
        result = ''
        for el in latex_code:
            result += '$' + el + '$' + '\n'

        fig, ax = plt.subplots(figsize=(4, 2))
        ax.text(0.5, 0.5, result, fontsize=20, ha='center', va='center')
        ax.set_axis_off()

        # Сохранение в буфер
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=dpi,
                    bbox_inches='tight', pad_inches=0.1)
        plt.close(fig)

        # Открытие изображения с помощью PIL
        buffer.seek(0)
        image = Image.open(buffer)

        if image.mode == 'RGBA':
            image = image.convert('RGB')
        # Сохраняем временное изображение
        base_path = PATH_TO_TEMP_FILES.joinpath(
            str(external_id)).joinpath(f'solution{random.randint(100,1000)}.jpg')
        base_path.parent.mkdir(parents=True, exist_ok=True)
        image.save(base_path, format='JPEG')  # Сохраняем как JPEG
        plt.close(fig)

        # # Инвертируем изображение
        # inverted_path = PATH_TO_TEMP_FILES.joinpath(
        #     str(external_id)).joinpath('solution_inverted.jpg')
        # BotService.invert_image(base_image_path=base_path,
        #                         output_path=inverted_path)

        # # Путь к финальному изображению с водяным знаком
        # final_path = PATH_TO_TEMP_FILES.joinpath(
        #     str(external_id)).joinpath('solution_with_watermark.jpg')

        # # Путь к файлу водяного знака

        # # Наложение водяного знака
        # BotService.add_watermark(base_image_path=inverted_path,
        #                          watermark_path=PATH_TO_TEMP_WATERMARK, output_path=final_path)

        return FSInputFile(path=base_path, filename="solution_with_watermark.jpg")

    def add_watermark(base_image_path, watermark_path, output_path):
        """
        Накладывает водяной знак с сохранением пропорций и регулируемой прозрачностью.
        """
        from PIL import ImageEnhance

        # Открываем базовое изображение и водяной знак
        base_image = Image.open(base_image_path).convert("RGBA")
        watermark = Image.open(watermark_path).convert("RGBA")

        # Масштабируем водяной знак с учётом `contain`
        base_width, base_height = base_image.size
        watermark_ratio = min(base_width / watermark.width,
                              base_height / watermark.height)
        watermark_width = int(watermark.width * watermark_ratio)
        watermark_height = int(watermark.height * watermark_ratio)
        watermark = watermark.resize(
            (watermark_width, watermark_height), Image.Resampling.LANCZOS)

        # Устанавливаем прозрачность водяного знака
        alpha = watermark.split()[3]  # Канал альфа (прозрачность)
        alpha = ImageEnhance.Brightness(alpha).enhance(0.5)  # Прозрачность 15%
        watermark.putalpha(alpha)

        # Позиция для водяного знака (по центру)
        position = (
            (base_width - watermark_width) // 2,
            (base_height - watermark_height) // 2
        )

        # Создаём слой для объединения изображений
        transparent = Image.new("RGBA", base_image.size, (0, 0, 0, 0))
        transparent.paste(base_image, (0, 0))
        transparent.paste(watermark, position, watermark)

        # Сохраняем итоговое изображение
        # Преобразуем в RGB для сохранения в JPEG
        transparent = transparent.convert("RGB")
        transparent.save(output_path, "JPEG")

    @staticmethod
    def create_image_by_user_requset(result: ChartResponse, external_id: int):
        plt.clf()

    # Отключаем интерактивный режим, чтобы не открывалось окно
        plt.ioff()
        x = eval(result.x_range)
        y = eval(result.y_range)
        plt.plot(x, y, label=result.title)

        plt.xlabel('x')
        plt.ylabel('y')
        plt.title(result.title)

        plt.legend()
        plt.grid(True)

        end_path = PATH_TO_TEMP_FILES.joinpath(
            str(external_id)).joinpath('chart.jpg')
        end_path.parent.mkdir(parents=True, exist_ok=True)
        plt.savefig(end_path, bbox_inches='tight', pad_inches=0.5, dpi=300)
        plt.close()
        return FSInputFile(path=end_path, filename="chart.jpg")

    @staticmethod
    async def check_user_subscription(event: Union[types.Message, types.CallbackQuery]):
        from Keyboards.keyboards import Keyboard
        try:
            member = await event.bot.get_chat_member('@student_helper_news', event.from_user.id)
            if member.status not in ["member", "creator", "administrator"]:
                await event.bot.send_message(
                    event.from_user.id,
                    text=LocalizationService.BotTexts.GetSubscriptionRequirements(
                        'ru'),
                    reply_markup=Keyboard.Get_Link_To_Channel('ru'),
                )
                return False
            return True
        except Exception as e:
            print(e)
            return False

    @staticmethod
    def calculate_remaining_time(subscription):
        if not subscription:
            return {
                "days": 0,
                "hours": 0,
                "minutes": 0
            }
        expiration_time = datetime.strptime(
            subscription.get('till'), '%Y-%m-%dT%H:%M:%S.%f')

        now = datetime.now()

        remaining_time = expiration_time - now

        if remaining_time.total_seconds() > 0:
            # Получаем количество дней и часов из оставшегося времени
            remaining_days = remaining_time.days
            remaining_seconds = remaining_time.seconds
            remaining_hours = remaining_seconds // 3600
            remaining_minutes = (remaining_seconds % 3600) // 60

            return {
                "days": remaining_days,
                "hours": remaining_hours,
                "minutes": remaining_minutes
            }
        else:
            return {
                "days": 0,
                "hours": 0,
                "minutes": 0
            }

    @staticmethod
    def get_my_profile_text(user, user_obj, subscription, selected_language):
        def format_limit(limit):
            if limit == 99999:
                return LocalizationService.BotTexts.GetUnlimitedTranslation(selected_language)
            elif limit == 0:
                return 0
            return limit

        localized_name = LocalizationService.BotTexts.GetsubscriptionName(
            name=subscription['sub_type']['name'] if subscription else 'inactive',
            selected_language=selected_language
        )
        remained_time = BotService.calculate_remaining_time(subscription)
        first_name = user.first_name

        if subscription:

            limitation_object = TelegramUserSubscriptionService.GetUserLimitations(
                user.id)
            limits = SUBSCRIPTION_LIMITATIONS
        else:
            limitation_object = TelegramUserSubscriptionService.GetUserDailyLimitations(
                user.id)
            limits = DAILY_LIMITATIONS

        limitations_remaining = limitation_object.get('limitations')

        localized_text = LocalizationService.BotTexts.GetMyProfileText(
            selected_language)

        return localized_text.format(
            first_name=first_name,
            user_id=user.id,
            localized_name=localized_name,
            days=remained_time['days'],
            hours=remained_time['hours'],
            minutes=remained_time['minutes'],
            default_mode_remain=limits["default_mode"] -
            max(0, limitations_remaining["default_mode"]),
            code_helper_remain=limits["code_helper"] -
            max(0, limitations_remaining["code_helper"]),
            abstract_writer_remain=limits["abstract_writer"] -
            max(0, limitations_remaining["abstract_writer"]),
            course_work_helper_remain=limits["course_work_helper"] -
            max(0, limitations_remaining["course_work_helper"]),
            essay_helper_remain=limits["essay_helper"] -
            max(0, limitations_remaining["essay_helper"]),
            photo_issue_helper_remain=limits["photo_issue_helper"] -
            max(0, limitations_remaining["photo_issue_helper"]),
            chart_creator_helper_remain=limits["chart_creator_helper"] - max(
                0, limitations_remaining["chart_creator_helper"]),
            power_point_helper_remain=limits["power_point_helper"] -
            max(0, limitations_remaining["power_point_helper"]),
            rewriting_helper_remain=limits["rewriting_helper"] -
            max(0, limitations_remaining["rewriting_helper"]),
            antiplagiat_helper_remain=limits['antiplagiat_helper'] -
            max(0, limitations_remaining['antiplagiat_helper'])
        )

    @staticmethod
    def parse_work_plan(plan_response: PlanResponse) -> str:
        result = []

        for heading in plan_response.headings:
            text = heading.heading_text.strip()

            if text.startswith("<h1>") and text.endswith("</h1>"):
                cleaned_text = re.sub(r"<h1>(.*?)</h1>", r"\1", text)
                result.append(f"• <b>{cleaned_text}</b>")
            elif text.startswith("<h2>") and text.endswith("</h2>"):
                cleaned_text = re.sub(r"<h2>(.*?)</h2>", r"\1", text)
                result.append(f"•• {cleaned_text}")

        return "\n".join(result)

    @staticmethod
    def getChangeLanguageText():
        ...

    @staticmethod
    async def GetTXTFileContent(bot: Bot, message: types.Message):
        document = message.document
        file = await bot.get_file(document.file_id)
        photo_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        photo_folder.mkdir(parents=True, exist_ok=True)
        file_path = photo_folder.joinpath(f'{file.file_unique_id}.txt')
        await bot.download(file=document.file_id, destination=file_path)
        content = ""
        async with aiofiles.open(file_path, 'r') as f:
            content = await f.read()
        return content

    @staticmethod
    async def WriteFileToTXT(content, external_id):
        end_path = PATH_TO_TEMP_FILES.joinpath(
            str(external_id)).joinpath(f'result{random.randint(1,10000)}.txt')
        async with aiofiles.open(end_path, 'w') as f:
            await f.write(content)
        return FSInputFile(path=end_path, filename="result.txt")

    @staticmethod
    async def WriteFileToDOCX(content, external_id):
        end_path = PATH_TO_TEMP_FILES.joinpath(
            str(external_id)).joinpath(f'result{random.randint(1,10000)}.docx')
        end_path.parent.mkdir(parents=True, exist_ok=True)
        doc = docx.Document()
        doc.add_paragraph(content)
        doc.save(end_path)
        return FSInputFile(path=end_path, filename="result.docx")

    @staticmethod
    def GetPPTXSettings(selected_language: str,
                        **kwargs):
        plain_text = kwargs.get("plain_text")
        language = kwargs.get("language")
        length = kwargs.get("length")
        template = kwargs.get("template")
        tone = kwargs.get("tone")
        fetch_images = kwargs.get("fetch_images")
        verbosity = kwargs.get("verbosity")

        base_settings = LocalizationService.BotTexts.GetPPTXSettings(
            selected_language)
        not_specified = LocalizationService.BotTexts.GetNotSpecified(
            selected_language)
        return base_settings.format(
            plain_text=plain_text if plain_text else not_specified,
            language=LocalizationService.BotTexts.GetHumanReadableLanguage(
                language) if language else not_specified,
            length=length if length else not_specified,
            template=template if template else not_specified,
            tone=LocalizationService.BotTexts.GetPPTXToneHR(
                tone, selected_language) if tone else not_specified,
            fetch_images=fetch_images if fetch_images else not_specified,
            verbosity=LocalizationService.BotTexts.GetPPTXVerbosityHR(
                verbosity, selected_language) if verbosity else not_specified
        )

    @staticmethod
    async def run_with_progress(message, task, total_time, *args, **kwargs):
        """
        Метод для отображения прогресс-бара с заданным временем, выполнения задачи и удаления сообщения.

        :param message: Сообщение, куда отправлять прогресс
        :param task: Асинхронная задача, которую нужно выполнить
        :param total_time: Общее время выполнения в секундах
        :param args: Аргументы для задачи
        :param kwargs: Именованные аргументы для задачи
        """
        progress_indicators = ["◒", "◐", "◓", "◑"]
        progress_cycle = itertools.cycle(progress_indicators)
        total_steps = 10  # Количество шагов для прогресс-бара
        step_time = total_time / total_steps  # Время на один шаг
        start_time = time.time()

        # Отправляем начальное сообщение
        progress_message = await message.answer("Прогресс начат...")

        for i in range(total_steps + 1):
            # Вычисляем оставшееся время
            remaining_time = max(total_time - (time.time() - start_time), 0)

            # Генерируем прогресс-бар и анимацию
            progress_bar = BotService.create_progress_bar(i, total_steps)
            animated_indicator = next(progress_cycle)

            # Обновляем текст сообщения
            await progress_message.edit_text(
                f"{progress_bar} {animated_indicator} (Осталось: {remaining_time:.1f} сек)"
            )

            # Задержка на заданный интервал
            if remaining_time > 0:
                await asyncio.sleep(step_time)

        # Выполняем задачу
        result = await task(*args, **kwargs)

        # Удаляем сообщение с прогресс-баром
        await progress_message.delete()

        # Возвращаем результат выполнения задачи
        return result

    @staticmethod
    def create_progress_bar(progress, total=10):
        """Создает строку с индикатором прогресса (эмодзи)."""
        filled_char = '⬛️'  # Черный квадрат
        empty_char = '⬜️'   # Белый квадрат
        filled_length = int(progress / total * total)
        empty_length = total - filled_length
        return filled_char * filled_length + empty_char * empty_length + f" {int(progress / total * 100)}%"

    @staticmethod
    async def send_long_message(target, text: str, parse_mode: str = None, disable_web_page_preview: bool = False, reply_markup=None):
        """
        Отправляет длинное сообщение в несколько батчей, поддерживая как Message, так и CallbackQuery.

        :param target: Объект типа `types.Message` или `types.CallbackQuery`.
        :param text: Текст сообщения.
        :param parse_mode: Режим разметки текста (например, Markdown или HTML).
        :param disable_web_page_preview: Отключение предпросмотра веб-страниц.
        """
        max_length = 4096  # Максимальная длина сообщения в Telegram

        # Функция для поиска незакрытых тегов
        def fix_unclosed_tags(batch, remaining_text):
            open_tags = re.findall(r'<([a-zA-Z]+)(?: [^>]*)?(?<!/)>', batch)
            close_tags = re.findall(r'</([a-zA-Z]+)>', batch)

            # Определяем незакрытые теги
            for tag in reversed(open_tags):
                if close_tags.count(tag) < open_tags.count(tag):
                    batch += f"</{tag}>"
                    remaining_text = f"<{tag}>" + remaining_text
                    open_tags.remove(tag)

            return batch, remaining_text

        # Разбиваем текст на батчи
        text_batches = [text[i:i + max_length]
                        for i in range(0, len(text), max_length)]

        # Определяем метод отправки в зависимости от типа объекта
        send_method = (
            target.reply if isinstance(
                target, types.Message) else target.message.reply
        )

        remaining_text = ""
        # Отправляем каждую часть сообщения
        for batch in text_batches:
            batch = remaining_text + batch
            batch, remaining_text = fix_unclosed_tags(batch, remaining_text)
            await send_method(
                text=batch,
                parse_mode=parse_mode,
                disable_web_page_preview=disable_web_page_preview,
                reply_markup=reply_markup,
            )

    @staticmethod
    def formatReferals(referal_obj: list, period: str):
        translation = {
            'last_month': 'в прошлом месяце',
            'this_month': 'в этом месяце',
            'all_time': 'за все время'
        }
        message = f"Топ рефералов{translation[period]} :\n"
        for i, referral in enumerate(referal_obj, start=1):
            username = referral.get('referal__username', 'Не указан')
            external_id = referral.get('referal__external_id', 'Не указан')
            invite_count = referral.get('invite_count', 0)
            message += f"{i}. Имя: {username}, ID: {external_id}, Кол-во приглашенных: {invite_count}\n"
        if len(referal_obj) == 0:
            message += 'Данных нет'
        return message

    @staticmethod
    async def getFileContent(bot: Bot, message: types.Message):
        document = message.document
        file = await bot.get_file(document.file_id)

        # Получаем расширение файла из имени
        file_extension = document.file_name.split(
            '.')[-1] if '.' in document.file_name else 'txt'

        # Создаем папку для пользователя, если её нет
        file_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        file_folder.mkdir(parents=True, exist_ok=True)

        # Формируем полный путь с правильным расширением
        file_path = file_folder.joinpath(
            f'{file.file_unique_id}.{file_extension}')
        await bot.download(file=document.file_id, destination=file_path)

        file_content = ''
        async with aiofiles.open(file_path, 'rb') as f:
            file_content = await f.read()
        return file_content

    @staticmethod
    async def GetWordFileContent(bot: Bot, message: types.Message):
        document = message.document
        file = await bot.get_file(document.file_id)
        file_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        file_folder.mkdir(parents=True, exist_ok=True)
        file_path = file_folder.joinpath(f'{file.file_unique_id}.docx')
        await bot.download(file=document.file_id, destination=file_path)
        doc = docx.Document(file_path)
        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        return content

    @staticmethod
    async def GetPDFFileContent(bot: Bot, message: types.Message):
        document = message.document
        file = await bot.get_file(document.file_id)
        file_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        file_folder.mkdir(parents=True, exist_ok=True)
        file_path = file_folder.joinpath(f'{file.file_unique_id}.pdf')
        await bot.download(file=document.file_id, destination=file_path)

        # Чтение содержимого PDF файла
        reader = PdfReader(file_path)
        content = ""
        for page in reader.pages:
            content += page.extract_text() + "\n"

        return content



    @staticmethod
    async def GetExcelFileContentJSON(bot: Bot, message: types.Message):
        document = message.document
        file = await bot.get_file(document.file_id)
        
        # Создание папки для сохранения файлов
        file_folder = PATH_TO_DOWNLOADED_FILES.joinpath(str(message.from_user.id))
        file_folder.mkdir(parents=True, exist_ok=True)
        file_path = file_folder.joinpath(f'{file.file_unique_id}.{document.file_name.split(".")[-1]}')
        await bot.download(file=document.file_id, destination=file_path)

        # Подготовка структуры данных для JSON
        data = {"sheets": []}

        # Проверяем тип файла и обрабатываем соответственно
        if file_path.suffix == '.xlsx':
            # Работа с .xlsx файлом через openpyxl
            workbook = load_workbook(file_path)
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                sheet_data = {
                    "name": sheet_name,
                    "rows": [
                        [str(cell) if cell is not None else "" for cell in row]  # Приведение всех данных к строковому типу
                        for row in worksheet.iter_rows(values_only=True)
                    ]
                }
                data["sheets"].append(sheet_data)
        elif file_path.suffix == '.xls':
            # Работа с .xls файлом через xlrd
            workbook = xlrd.open_workbook(file_path)
            for sheet_index in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sheet_index)
                sheet_data = {
                    "name": sheet.name,
                    "rows": [
                        [str(cell) if cell is not None else "" for cell in sheet.row_values(row_index)]  # Приведение всех данных к строковому типу
                        for row_index in range(sheet.nrows)
                    ]
                }
                data["sheets"].append(sheet_data)
        else:
            return json.dumps({"error": "Unsupported file format. Only .xls and .xlsx are allowed."}, ensure_ascii=False)

        # Возвращаем данные в формате JSON с сохранением русских символов
        return json.dumps(data, ensure_ascii=False, indent=4)





    @staticmethod
    async def run_process_with_countdown(
            message: types.Message = None,
            phrases=["Анализирую запрос...",
                     "Ищу информацию...", "Генерирую текст...", ],
            task=None, *args, **kwargs
            ):
        """
        Метод для отображения таймера, смены фраз и выполнения задачи с удалением сообщения по завершению.

        :param message: Сообщение, куда отправлять прогресс
        :param task: Асинхронная задача, которую нужно выполнить
        :param total_time: Общее время выполнения в секундах
        :param args: Аргументы для задачи
        :param kwargs: Именованные аргументы для задачи
        """

        phrase_cycle = itertools.cycle(phrases)
        start_time = time.time()

        # Отправляем начальное сообщение

        current_phrase = next(phrase_cycle)  # Первая фраза

        # Отправляем начальное сообщение
        progress_message = await message.answer(f"<code>[00:00] {current_phrase}</code>")

        current_time = 0  # Отслеживание времени
        phrase_update_time = 0  # Для контроля смены фразы

        previous_message_content = f"<code>[00:00] {current_phrase}</code>"

        task_coro = task(*args, **kwargs)
        task_result = asyncio.create_task(task_coro)

        while not task_result.done():
            # Обновляем таймер каждую секунду
            current_time = int(time.time() - start_time)
            minutes, seconds = divmod(current_time, 60)

            # Обновляем фразу каждые две секунды
            if current_time - phrase_update_time >= 2:
                current_phrase = next(phrase_cycle)
                phrase_update_time = current_time

            # Генерируем новое сообщение
            new_message_content = f"<code>[{minutes:02}:{seconds:02}] {current_phrase}</code>"

            # Проверяем, изменилось ли содержимое
            if new_message_content != previous_message_content:
                await progress_message.edit_text(new_message_content)
                previous_message_content = new_message_content

            # Ждем 1 секунду перед следующей итерацией
            await asyncio.sleep(1)

        result = await task_result

        await progress_message.delete()
        return result

    def __escape_text(text):
        """
        Экранирует символы &, <, >
        """
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def __process_code_block(segment):
        """
        Если сегмент – разрешённый блок кода (начинается с <pre> или <code>),
        то оставляем сами теги, а содержимое между ними обрабатываем через escape_text.
        Учитываем следующие варианты:
        1. <pre><code class="…"> ... </code></pre>
        2. <code class="…"> ... </code>
        3. <pre …> ... </pre>
        4. <code> ... </code>
        Если ничего не подошло, возвращаем сегмент без изменений.
        """
        patterns = [
            # 1. Обрабатываем конструкцию <pre><code …>...</code></pre>
            (r'^(<pre><code[^>]*>)(.*?)(</code></pre>)$', re.DOTALL),
            # 2. Блок <code class="…">...</code>
            (r'^(<code class="[^">]*">)(.*?)(</code>)$', re.DOTALL),
            # 3. Блок <pre …>...</pre>
            (r'^(<pre[^>]*>)(.*?)(</pre>)$', re.DOTALL),
            # 4. Блок <code>...</code>
            (r'^(<code>)(.*?)(</code>)$', re.DOTALL)
        ]
        for pattern, flags in patterns:
            m = re.match(pattern, segment, flags)
            if m:
                opening, inner, closing = m.groups()
                # Экранируем содержимое внутри кода
                return opening + BotService.__escape_text(inner) + closing
        return segment

    def escape_html(input_text):
        """
        Метод экранирования:
        – Все символы &, <, > заменяются на сущности, если они находятся вне разрешённых блоков.
        – Разрешённые блоки (совпадающие по одному из шаблонов ниже) остаются целиком,
            за исключением блоков кода (<pre>, <code>), в которых заменяем спецсимволы только внутри контента.

        Разрешённые шаблоны (эти фрагменты не должны изменяться, за исключением обработки кода):
        1. <pre><code class=".*?">.*?</code></pre>
        2. <code class=".*?">.*?</code>
        3. <pre.*?>.*?</pre>
        4. <code>.*?</code>
        5. <b>.*?</b>
        6. <i>.*?</i>
        7. <u>.*?</u>
        8. <s>.*?</s>
        9. <span class="tg-spoiler">.*?</span>
        10. <a href=".*?">.*?</a>
        """
        # Список разрешённых шаблонов – порядок имеет значение, более "специфичные" (блоки кода) ставим первыми.
        allowed_patterns = [
            r'<pre><code class=".*?">.*?</code></pre>',
            r'<code class=".*?">.*?</code>',
            r'<pre.*?>.*?</pre>',
            r'<code>.*?</code>',
            r'<b>.*?</b>',
            r'<i>.*?</i>',
            r'<u>.*?</u>',
            r'<s>.*?</s>',
            r'<span class="tg-spoiler">.*?</span>',
            r'<a href=".*?">.*?</a>'
        ]
        # Объединяем их в одно регулярное выражение через ОR.
        combined_pattern = "(" + "|".join(allowed_patterns) + ")"
        allowed_regex = re.compile(combined_pattern, flags=re.DOTALL)

        result_segments = []
        last_index = 0

        # Ищем все вхождения разрешённых блоков
        for match in allowed_regex.finditer(input_text):
            start, end = match.span()
            # Текст между разрешёнными блоками — экранируем полностью.
            non_allowed = input_text[last_index:start]
            result_segments.append(BotService.__escape_text(non_allowed))

            allowed_segment = match.group(0)
            # Если блок начинается с <pre или <code>, то считаем его кодом,
            # обрабатываем содержимое с escape_text (но оставляем сами теги)
            seg_strip = allowed_segment.lstrip()
            if seg_strip.startswith("<pre") or seg_strip.startswith("<code"):
                result_segments.append(
                    BotService.__process_code_block(allowed_segment))
            else:
                # Остальные разрешённые блоки оставляем нетронутыми.
                result_segments.append(allowed_segment)
            last_index = end

        # Экранируем остаток текста после последнего совпадения.
        result_segments.append(
            BotService.__escape_text(input_text[last_index:]))

        return "".join(result_segments)

    @staticmethod
    async def go_menu(bot: Bot, event: Union[Message, CallbackQuery], state: FSMContext, final_state: State):
        from Keyboards import Keyboard

        if isinstance(event, CallbackQuery):
            message = event.message
            user_id = event.from_user.id
        else:
            message = event
            user_id = message.from_user.id

        try:
            await bot.unpin_all_chat_messages(chat_id=message.chat.id)
        except Exception:
            pass

        # Достаём язык и клавиатуру
        data = await state.get_data()
        language = data.get('language', 'ru')

        text = LocalizationService.BotTexts.GetInstrumentsText(language)
        markup = Keyboard.Get_Instruments(user_id, language)

        try:
            if message.text:
                await message.edit_text(text=text, reply_markup=markup)
            else:
                raise ValueError("Cannot edit non-text message")
        except Exception:
            try:
                await message.delete()
            except Exception:
                ...
            await bot.send_message(chat_id=message.chat.id, text=text, reply_markup=markup)

        # Устанавливаем новое состояние
        await state.set_state(final_state)

    @staticmethod
    def GetSubscriptionPrice(selected_language, subscription):
        subscription_text = LocalizationService.BotTexts.SubscriptionText(
            selected_language)
        return subscription_text.format(price=subscription.get('price'))

    @staticmethod
    async def download_file(bot: Bot, message: types.Message):
        document = message.document

        file = await bot.get_file(document.file_id)

        photo_folder = PATH_TO_DOWNLOADED_FILES.joinpath(
            str(message.from_user.id))
        photo_folder.mkdir(parents=True, exist_ok=True)

        # Get the file name from the document object
        original_file_name = document.file_name
        file_extension = original_file_name.split(
            '.')[-1] if '.' in original_file_name else ''

        file_path = photo_folder.joinpath(
            f'{file.file_unique_id}.{file_extension}')
        await bot.download(file=document.file_id, destination=file_path)

        return file_path
